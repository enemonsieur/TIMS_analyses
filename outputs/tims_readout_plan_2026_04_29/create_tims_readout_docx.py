from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


output_directory = Path(
    r"C:\Users\njeuk\OneDrive\Documents\Charite Berlin\TIMS\outputs\tims_readout_plan_2026_04_29"
)
output_directory.mkdir(parents=True, exist_ok=True)
document_path = output_directory / "One-Month_TIMS_Readout_Plan.docx"


NAVY = "1F4E5F"
TEAL = "0F766E"
LIGHT_TEAL = "E6F4F1"
LIGHT_GRAY = "F4F6F7"
DARK = "1F2933"
MUTED = "5B6670"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9DEE3", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=9.5, bold=False, color=DARK):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(document, text="", size=9.4, color=DARK, bold=False, before=0, after=3, line=1.0):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    return paragraph


def add_section_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text.upper())
    style_run(run, size=8.4, bold=True, color=TEAL)
    return paragraph


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            set_cell_width(row.cells[idx], width)


document = Document()
section = document.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.42)
section.left_margin = Inches(0.52)
section.right_margin = Inches(0.52)

styles = document.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
styles["Normal"].font.size = Pt(9.4)

title = document.add_paragraph()
title.paragraph_format.space_after = Pt(0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_run = title.add_run("One-Month TIMS Readout Plan")
style_run(title_run, size=20, bold=True, color=NAVY)

subtitle = document.add_paragraph()
subtitle.paragraph_format.space_before = Pt(1)
subtitle.paragraph_format.space_after = Pt(7)
subtitle_run = subtitle.add_run("Prepared for Francesca | Start date: April 29, 2026 | Weekly touchpoint: Francesca, Khalid, and N")
style_run(subtitle_run, size=8.8, color=MUTED)

add_section_heading(document, "Objective")
add_paragraph(
    document,
    "I want to keep the next month focused. I would not add MEG into this first month, because that would add a separate lab, booking, acquisition setup, and analysis pipeline. The first step is simpler: show whether TIMS gives us a recoverable EEG response, and whether that response changes with stimulation intensity in a way that is comparable to the logic we use in TMS.",
)
add_paragraph(
    document,
    "Success does not require every metric to work. The month is successful if one or two readouts are reliable enough to justify the next step: a TIMS-evoked potential, ITPC/PLV or coherence, resting-state power or delta stability, entropy, or cortical excitability before versus after stimulation.",
)

add_section_heading(document, "Protocol Logic")
add_paragraph(
    document,
    "We start with the 100% theta-burst stimulation protocol on Khalid, focused on M1 first because it gives the clearest comparison point to TMS-EEG and MEP-based excitability. We recover the immediate event-related response, then test whether the signal is clean enough to interpret and not dominated by artifact.",
)
add_paragraph(
    document,
    "In parallel, we run before/after readouts: resting-state EEG for power, delta stability, and entropy, plus sparse pulse stimulation to test whether cortical excitability changes after TIMS. If the 100% protocol gives a convincing signal, we move into dose-response testing at 100%, 50%, 25%, and 10% using the same acquisition and analysis logic.",
)

add_section_heading(document, "Experiment Structure")
diagram = document.add_table(rows=1, cols=7)
diagram.alignment = WD_TABLE_ALIGNMENT.CENTER
diagram.autofit = False
diagram_widths = [1.85, 0.3, 1.85, 0.3, 1.85, 0.3, 1.85]
set_table_width(diagram, diagram_widths)
diagram_steps = [
    ("1. 100% M1 TIMS", "Khalid; recover evoked response + before/after EEG"),
    (">", ""),
    ("2. Metric decision", "TEP first; otherwise ITPC/PLV, power, entropy, excitability"),
    (">", ""),
    ("3. Dose response", "Repeat at 100%, 50%, 25%, and 10%"),
    (">", ""),
    ("4. Replication", "Four additional participants; individual + group readout"),
]
for idx, (label, detail) in enumerate(diagram_steps):
    cell = diagram.cell(0, idx)
    set_cell_width(cell, diagram_widths[idx])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if detail:
        set_cell_shading(cell, LIGHT_TEAL)
        set_cell_border(cell, color="B9D8D3")
        set_cell_margins(cell, top=90, bottom=90, start=80, end=80)
        title_run = paragraph.add_run(label)
        style_run(title_run, size=7.6, bold=True, color=NAVY)
        paragraph.add_run("\n")
        detail_run = paragraph.add_run(detail)
        style_run(detail_run, size=6.7, color=DARK)
    else:
        set_cell_shading(cell, "FFFFFF")
        set_cell_border(cell, color="FFFFFF")
        set_cell_margins(cell, top=80, bottom=80, start=25, end=25)
        arrow_run = paragraph.add_run(label)
        style_run(arrow_run, size=9.0, bold=True, color=TEAL)

add_section_heading(document, "Monthly Layout")
table = document.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
headers = ["Week", "Key action", "Deliverable", "Main constraint", "People"]
widths = [1.0, 2.05, 2.45, 1.65, 1.15]
set_table_width(table, widths)

for idx, header in enumerate(headers):
    cell = table.rows[0].cells[idx]
    set_cell_shading(cell, NAVY)
    set_cell_border(cell, color=NAVY)
    set_cell_margins(cell, top=90, bottom=90, start=85, end=85)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(header)
    style_run(run, size=7.9, bold=True, color="FFFFFF")

rows = [
    (
        "Week 1\nApr 29-May 3",
        "Finalize the analysis pipeline, run the field/depth simulation from the core drive, and recover first TIMS-evoked responses from available data.",
        "Short simulation output plus first EEG plots showing whether we can recover a clean response and phase-consistency metrics.",
        "Pipeline sorting and true response versus stimulation artifact.",
        "Khalid + N",
    ),
    (
        "Week 2\nMay 4-May 10",
        "Run the 100% M1 theta-burst protocol on Khalid and analyze immediate, resting-state, and excitability readouts.",
        "First decision memo: best metric, whether the response is TMS-like enough to compare, and whether to move into dose response.",
        "Not all metrics will be reliable; pivot to the strongest readout.",
        "Khalid + N; Francesca review",
    ),
    (
        "Week 3\nMay 11-May 17",
        "Repeat the same protocol across 100%, 50%, 25%, and 10% stimulation intensity.",
        "Preliminary dose-response trend across the strongest metrics: TIMS-evoked response, ITPC/PLV/coherence, power/entropy, or excitability.",
        "Order effects, fatigue, and keeping the protocol identical.",
        "Khalid + N",
    ),
    (
        "Week 4\nMay 18-May 24",
        "Replicate the most promising protocol in four additional participants or patients.",
        "Individual plots plus a first group-average summary showing whether the effect replicates beyond N=1.",
        "Scheduling and inter-individual variability.",
        "Khalid + N; Francesca unblockers",
    ),
]

for row_idx, values in enumerate(rows, start=1):
    cells = table.add_row().cells
    for col_idx, value in enumerate(values):
        cell = cells[col_idx]
        set_cell_width(cell, widths[col_idx])
        set_cell_shading(cell, LIGHT_TEAL if row_idx % 2 else "FFFFFF")
        set_cell_border(cell)
        set_cell_margins(cell, top=105, bottom=105, start=85, end=85)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(value)
        style_run(run, size=7.35, bold=(col_idx == 0), color=DARK)

add_section_heading(document, "Decision Rule")
add_paragraph(
    document,
    "By the end of Week 2, we choose the metric that carries the project forward. If we have a clean TIMS-evoked potential, that becomes the anchor. If not, we move with the strongest reliable alternative: phase consistency, resting-state power or delta stability, entropy, or cortical excitability. By the end of Week 4, the deliverable is a concise result package showing whether TIMS has a plausible physiological effect, whether it scales with intensity, and whether it replicates across participants.",
    after=2,
)

note_table = document.add_table(rows=1, cols=1)
note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
note_table.autofit = False
note_cell = note_table.cell(0, 0)
set_cell_width(note_cell, 8.35)
set_cell_shading(note_cell, LIGHT_GRAY)
set_cell_border(note_cell, color="E1E6EA")
set_cell_margins(note_cell, top=95, bottom=95, start=120, end=120)
note_paragraph = note_cell.paragraphs[0]
note_paragraph.paragraph_format.space_after = Pt(0)
note_run = note_paragraph.add_run(
    "Metric language is aligned with standard TMS-EEG readouts: evoked potentials, GMFA/GMFP, time-frequency measures, phase consistency, and artifact constraints; plus recent descriptions of ITPC and canonical TEP windows."
)
style_run(note_run, size=7.6, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("TIMS planning memo | April 2026")
style_run(footer_run, size=7.5, color=MUTED)

document.core_properties.author = "TIMS"
document.core_properties.title = "One-Month TIMS Readout Plan"
document.core_properties.subject = "TIMS EEG readout planning"
document.save(document_path)
print(document_path)
